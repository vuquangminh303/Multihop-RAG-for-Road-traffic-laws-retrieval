import os
import re
import glob
import json
import sys  # Añadido para aumentar el límite de recursión
from docx import Document
from pdfminer.high_level import extract_text as pdfminer_extract_text
import pytesseract
from pdf2image import convert_from_path

# Aumentar el límite de recursión
sys.setrecursionlimit(10000)

# Kiểm tra PDF có văn bản chọn được hay không
def has_selectable_text(pdf_path):
    try:
        text = pdfminer_extract_text(pdf_path)
        return len(text.strip()) > 0
    except Exception as e:
        print(f"Error checking selectable text in {pdf_path}: {e}")
        return False

# Trích xuất văn bản từ PDF bằng OCR
def extract_text_from_pdf_with_ocr(pdf_path):
    try:
        images = convert_from_path(pdf_path)
        text = ""
        for image in images:
            text += pytesseract.image_to_string(image, lang="vie")
        return text
    except Exception as e:
        print(f"OCR error for {pdf_path}: {e}")
        return ""

# Trích xuất văn bản từ PDF
def extract_text_from_pdf(pdf_path):
    try:
        if has_selectable_text(pdf_path):
            return pdfminer_extract_text(pdf_path)
        else:
            return extract_text_from_pdf_with_ocr(pdf_path)
    except RecursionError:
        print(f"Recursion error in PDF processing for {pdf_path}. Trying alternative method.")
        try:
            # Método alternativo para PDFs problemáticos - intentar OCR directamente
            return extract_text_from_pdf_with_ocr(pdf_path)
        except Exception as e:
            print(f"Alternative method failed for {pdf_path}: {e}")
            return ""

# Trích xuất văn bản từ DOCX (bao gồm bảng)
def extract_text_from_docx(file_path):
    doc = Document(file_path)
    
    # Lấy text từ paragraphs
    paragraphs_text = [para.text for para in doc.paragraphs]
    
    # Lấy text từ tables
    tables_text = []
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells]
            tables_text.append(" | ".join(row_text))
    
    # Kết hợp text từ paragraphs và tables
    all_text = "\n".join(paragraphs_text + tables_text)
    return all_text

# Trích xuất văn bản dựa trên loại file
def extract_text(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".docx":
        return extract_text_from_docx(file_path)
    elif ext == ".pdf":
        return extract_text_from_pdf(file_path)
    else:
        raise ValueError("Unsupported file type")

def extract_document_id(text):
    # Cải thiện regex để bắt nhiều dạng ID
    patterns = [
        r"Số:\s*(\d+/\d+/TT-\w+)",
        r"Số:\s*(\d+\.\d+\.TT\.[\w\.]+)",
        r"(TT\.\d+\.\d+\.[\w\.]+)",
        r"([A-Z]+\d+\.signed)",
        r"(\d+[_/]\d+[_/](?:TT|TTLT)[-_][\w]+)"
    ]
    
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            return match.group(1)
    
    # Nếu không tìm thấy, thử lấy từ tên file
    return "Unknown"

# Trích xuất tiêu đề
def extract_title(text):
    lines = text.split("\n")
    
    # Tìm dòng có "THÔNG TƯ"
    for i, line in enumerate(lines):
        if "THÔNG TƯ" in line.upper():
            # Tìm dòng tiếp theo có nội dung
            for j in range(i+1, min(i+5, len(lines))):
                if lines[j].strip() and "BỘ" not in lines[j] and "Căn cứ" not in lines[j]:
                    return lines[j].strip()
    
    return "Unknown"

# Trích xuất cơ quan ban hành
def extract_issuing_agency(text):
    patterns = [
        r"(BỘ [^\n]+)",
        r"(CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM[\s\S]*?)\n"
    ]
    
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            return match.group(1).strip()
    
    return "Unknown"

# Trích xuất ngày ban hành
def extract_date(text):
    patterns = [
        r"Hà Nội, (ngày \d+ tháng \d+ năm \d+)",
        r"ngày (\d+ tháng \d+ năm \d+)",
        r"(\d+/\d+/\d+)",
        r"(\d+ tháng \d+ năm \d+)"
    ]
    
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            return match.group(1).strip()
    
    return "Unknown"

# Phân đoạn văn bản theo "Điều X"
def split_into_segments(text):
    pattern = r"(Điều \d+\..*?)(?=Điều \d+\.|$)"
    segments = re.findall(pattern, text, flags=re.DOTALL)
    
    # Nếu không tìm thấy segments theo "Điều", thử phân đoạn theo các phần
    if not segments:
        pattern = r"(Chương \d+\..*?)(?=Chương \d+\.|$)"
        segments = re.findall(pattern, text, flags=re.DOTALL)
    
    # Nếu vẫn không tìm thấy, chia thành các đoạn nhỏ hơn
    if not segments and len(text) > 1000:
        segments = [text[i:i+1000] for i in range(0, len(text), 1000)]
    elif not segments:
        segments = [text]
    
    return [s.strip() for s in segments if s.strip()]


# Đường dẫn thư mục
directory = "/home/cong/workspace/chatbot-retrieval-based/Thongtu"
docx_files = glob.glob(os.path.join(directory, "*.docx"))
pdf_files = glob.glob(os.path.join(directory, "*.pdf"))
all_files = docx_files + pdf_files

segment_data = []
for file_path in all_files:
    try:
        print(f"Processing {file_path}...")
        text = extract_text(file_path)
        
        # Lấy filename để dùng làm fallback nếu không trích xuất được id
        filename = os.path.basename(file_path)
        doc_id = extract_document_id(text)
        
        # Sử dụng tên file nếu không trích xuất được ID
        if doc_id == "Unknown":
            doc_id = os.path.splitext(filename)[0]
        
        title = extract_title(text)
        agency = extract_issuing_agency(text)
        date = extract_date(text)
        segments = split_into_segments(text)
        
        # Kiểm tra và ghi log nếu không có segments
        if not segments:
            print(f"Warning: No segments found in {file_path}")
            segments = [text]
            
        for i, segment in enumerate(segments):
            segment_data.append({
                "document_id": doc_id,
                "title": title,
                "issuing_agency": agency,
                "date": date,
                "segment_id": f"{doc_id}_{i}",
                "content": segment
            })
        print(f"Successfully processed {file_path} with {len(segments)} segments")
    except Exception as e:
        print(f"Error processing {file_path}: {e}")

# Lưu metadata vào file JSON
with open("metadata.json", "w", encoding="utf-8") as f:
    json.dump(segment_data, f, ensure_ascii=False, indent=4)

print(f"Metadata saved to metadata.json with {len(segment_data)} segments from {len(all_files)} files")
print("\nĐể tạo vector embeddings, hãy chạy script create_embeddings.py tiếp theo.")